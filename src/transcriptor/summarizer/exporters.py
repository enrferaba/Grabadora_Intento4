"""Summary export helpers."""
from __future__ import annotations

import io
import json
from typing import Tuple

from docx import Document

from .engine import SummaryDocument


def export_markdown(document: SummaryDocument) -> bytes:
    lines = [f"# {document.title}"]
    if document.client:
        lines.append(f"**Cliente:** {document.client}")
    if document.meeting_date:
        lines.append(f"**Fecha:** {document.meeting_date}")
    if document.attendees:
        lines.append("**Asistentes:** " + ", ".join(document.attendees))
    lines.append("")
    lines.append(document.summary)
    lines.append("")
    lines.append("## Puntos clave")
    for point in document.key_points:
        lines.append(f"- {point}")
    lines.append("")
    lines.append("## Acciones")
    for action in document.actions:
        due = f" (para {action.due})" if action.due else ""
        lines.append(f"- **{action.owner}:** {action.task}{due}")
    lines.append("")
    if document.risks:
        lines.append("## Riesgos")
        for risk in document.risks:
            lines.append(f"- {risk}")
        lines.append("")
    if document.next_steps:
        lines.append("## Próximos pasos")
        for step in document.next_steps:
            lines.append(f"- {step}")
        lines.append("")
    if document.fallback_used:
        lines.append("_Resumen generado mediante modo extractivo por limitaciones del dispositivo o licencia._")
    return "\n".join(lines).encode("utf-8")


def export_json(document: SummaryDocument) -> bytes:
    payload = document.as_dict()
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")


def export_docx(document: SummaryDocument) -> bytes:
    doc = Document()
    doc.add_heading(document.title, level=1)
    if document.client or document.meeting_date:
        paragraph = doc.add_paragraph()
        if document.client:
            paragraph.add_run("Cliente: ").bold = True
            paragraph.add_run(document.client)
            paragraph.add_run("\n")
        if document.meeting_date:
            paragraph.add_run("Fecha: ").bold = True
            paragraph.add_run(document.meeting_date)
            paragraph.add_run("\n")
    if document.attendees:
        paragraph = doc.add_paragraph()
        paragraph.add_run("Asistentes: ").bold = True
        paragraph.add_run(", ".join(document.attendees))

    doc.add_heading("Resumen ejecutivo", level=2)
    doc.add_paragraph(document.summary)

    doc.add_heading("Puntos clave", level=2)
    for point in document.key_points:
        doc.add_paragraph(point, style="List Bullet")

    doc.add_heading("Acciones", level=2)
    for action in document.actions:
        description = f"{action.owner}: {action.task}"
        if action.due:
            description += f" (para {action.due})"
        doc.add_paragraph(description, style="List Bullet")

    if document.risks:
        doc.add_heading("Riesgos", level=2)
        for risk in document.risks:
            doc.add_paragraph(risk, style="List Bullet")

    if document.next_steps:
        doc.add_heading("Próximos pasos", level=2)
        for step in document.next_steps:
            doc.add_paragraph(step, style="List Bullet")

    if document.fallback_used:
        doc.add_paragraph("Resumen generado en modo extractivo por limitaciones detectadas.")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_document(document: SummaryDocument, export_format: str) -> Tuple[bytes, str]:
    if export_format == "markdown":
        return export_markdown(document), "text/markdown"
    if export_format == "json":
        return export_json(document), "application/json"
    if export_format == "docx":
        return export_docx(document), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    raise ValueError(f"Formato no soportado: {export_format}")

